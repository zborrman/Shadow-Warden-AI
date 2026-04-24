"""
warden/api/file_scan.py
───────────────────────
POST /filter/file — scan a file for PII, secrets, and prompt-injection
                    before it is uploaded to an external AI tool.

Target audience: Community Business (SMB) tier.
Supports: .txt, .md, .csv, .json, .py, .js, .ts, .html, .pdf (text layer)

Typical workflow
────────────────
  1. User drags a file onto a browser extension or desktop agent.
  2. Agent calls  POST /filter/file  with multipart/form-data.
  3. Warden extracts text, runs SecretRedactor + SemanticGuard.
  4. Response: { safe: bool, risk_level: str, findings: [...], sanitized_text: str }
  5. If safe=True the agent proceeds to upload to AI tool.
  6. If safe=False the agent blocks and shows the user a summary.

Risk levels: SAFE / LOW / MEDIUM / HIGH / CRITICAL
"""
from __future__ import annotations

import io
import logging
import time
from typing import Annotated

from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from pydantic import BaseModel

log = logging.getLogger("warden.api.file_scan")

router = APIRouter(prefix="/filter", tags=["file-scan"])

# ── Max file size ─────────────────────────────────────────────────────────────
_MAX_BYTES = 10 * 1024 * 1024  # 10 MB
_MAX_TEXT_CHARS = 200_000       # truncate extracted text for scanning

# ── Supported MIME / extension map ────────────────────────────────────────────
_TEXT_TYPES = {
    "text/plain", "text/markdown", "text/csv",
    "text/html", "text/xml",
    "application/json",
    "application/javascript",
    "application/typescript",
    "application/x-python", "text/x-python",
}


class FileScanFinding(BaseModel):
    kind:    str   # "secret" | "pii" | "prompt_injection" | "high_entropy"
    label:   str   # human-readable (e.g. "OpenAI API Key")
    excerpt: str   # redacted excerpt for display
    line:    int   # approximate line number (0 = unknown)


class FileScanResponse(BaseModel):
    filename:       str
    size_bytes:     int
    safe:           bool
    risk_level:     str   # SAFE / LOW / MEDIUM / HIGH / CRITICAL
    findings:       list[FileScanFinding]
    findings_count: int
    sanitized_text: str   # text with secrets replaced
    processing_ms:  float
    truncated:      bool


def _extract_text(data: bytes, content_type: str, filename: str) -> str:
    """Return plain text from uploaded bytes."""
    ct = (content_type or "").lower().split(";")[0].strip()
    fname = (filename or "").lower()

    if ct == "application/pdf" or fname.endswith(".pdf"):
        return _extract_pdf(data)

    if (ct == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or fname.endswith(".docx")):
        return _extract_docx(data)

    if (ct in ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
               "application/vnd.ms-excel")
            or fname.endswith((".xlsx", ".xls"))):
        return _extract_xlsx(data)

    try:
        return data.decode("utf-8", errors="replace")
    except Exception:
        return data.decode("latin-1", errors="replace")


def _extract_docx(data: bytes) -> str:
    """Extract paragraphs + table cells from a .docx file."""
    try:
        import docx  # python-docx
        doc = docx.Document(io.BytesIO(data))
        parts: list[str] = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(c.text for c in row.cells))
        return "\n".join(parts)
    except ImportError:
        log.debug("python-docx not installed — falling back to raw decode")
        return data.decode("utf-8", errors="replace")
    except Exception as exc:
        log.debug("docx extract failed: %s", exc)
        return data.decode("utf-8", errors="replace")


def _extract_xlsx(data: bytes) -> str:
    """Extract cell values from all sheets of a .xlsx file."""
    try:
        import openpyxl  # type: ignore
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        rows: list[str] = []
        for ws in wb.worksheets:
            rows.append(f"[Sheet: {ws.title}]")
            for row in ws.iter_rows(values_only=True):
                rows.append("\t".join("" if c is None else str(c) for c in row))
        wb.close()
        return "\n".join(rows)
    except ImportError:
        log.debug("openpyxl not installed — falling back to raw decode")
        return data.decode("utf-8", errors="replace")
    except Exception as exc:
        log.debug("xlsx extract failed: %s", exc)
        return data.decode("utf-8", errors="replace")


def _extract_pdf(data: bytes) -> str:
    """Extract text layer from PDF — fail-open with raw bytes decode."""
    try:
        from pdfminer.high_level import extract_text as pm_extract  # type: ignore
        return pm_extract(io.BytesIO(data))
    except ImportError:
        pass
    try:
        from pypdf import PdfReader  # type: ignore
        reader = PdfReader(io.BytesIO(data))
        return "\n".join(p.extract_text() or "" for p in reader.pages)
    except ImportError:
        pass
    # Last resort — treat as raw text (catches text-based PDFs)
    return data.decode("utf-8", errors="replace")


def _line_of_match(text: str, start: int) -> int:
    return text[:start].count("\n") + 1


def _risk_level(findings: list[FileScanFinding], injection_found: bool) -> str:
    if injection_found:
        return "CRITICAL"
    n = len(findings)
    if n == 0:
        return "SAFE"
    secrets = sum(1 for f in findings if f.kind == "secret")
    if secrets >= 3 or n >= 10:
        return "HIGH"
    if secrets >= 1 or n >= 4:
        return "MEDIUM"
    return "LOW"


@router.post("/file", response_model=FileScanResponse, summary="Scan file before AI upload")
async def scan_file(
    file: Annotated[UploadFile, File(description="File to scan (max 10 MB)")],
    tenant_id: Annotated[str, Form()] = "default",
    strict: Annotated[bool, Form()] = False,
) -> FileScanResponse:
    """
    Scan a file for secrets, PII, and prompt-injection before uploading to an AI tool.

    Returns a risk assessment and a sanitized version of the text with secrets redacted.
    Safe for use on personal computers — no file content is stored or logged.
    """
    t0 = time.perf_counter()

    # ── Size guard ────────────────────────────────────────────────────────────
    raw = await file.read(_MAX_BYTES + 1)
    if len(raw) > _MAX_BYTES:
        raise HTTPException(
            status_code=413,
            detail=f"File too large (max {_MAX_BYTES // 1024 // 1024} MB).",
        )

    size_bytes = len(raw)
    filename   = file.filename or "unknown"

    # ── Text extraction ───────────────────────────────────────────────────────
    text = _extract_text(raw, file.content_type or "", filename)
    truncated = len(text) > _MAX_TEXT_CHARS
    if truncated:
        text = text[:_MAX_TEXT_CHARS]

    # ── Secret + PII scan ─────────────────────────────────────────────────────
    from warden.schemas import RedactionPolicy
    from warden.secret_redactor import SecretRedactor

    redactor = SecretRedactor(strict=strict)
    result   = redactor.redact(text, RedactionPolicy.FULL)

    findings: list[FileScanFinding] = []
    for sf in result.findings:
        line = _line_of_match(text, getattr(sf, "start", 0))
        findings.append(FileScanFinding(
            kind    = "pii" if sf.kind in ("email", "ssn", "iban", "credit_card") else "secret",
            label   = sf.kind.replace("_", " ").title(),
            excerpt = sf.redacted_to[:80] if sf.redacted_to else f"[{sf.kind}]",
            line    = line,
        ))

    sanitized = result.text

    # ── Prompt-injection scan ─────────────────────────────────────────────────
    injection_found = False
    try:
        from warden.semantic_guard import RiskLevel, SemanticGuard
        guard = SemanticGuard()
        analysis = guard.analyse(sanitized[:5_000])  # fast scan on first 5k chars
        if analysis.risk_level in (RiskLevel.HIGH, RiskLevel.BLOCK):
            injection_found = True
            top = analysis.top_flag
            excerpt = (top.detail[:80] if top and top.detail else "suspicious pattern")
            findings.append(FileScanFinding(
                kind    = "prompt_injection",
                label   = "Potential Prompt Injection",
                excerpt = excerpt,
                line    = 0,
            ))
    except Exception as exc:
        log.debug("file_scan: semantic guard skipped: %s", exc)

    # ── Obfuscation check (base64 blobs, hex, ROT13) ─────────────────────────
    try:
        from warden.obfuscation import decode as obfuscation_decode
        result_ob = obfuscation_decode(sanitized[:10_000])
        depth = len(result_ob.layers_found)
        if depth >= 1 and result_ob.decoded_extra:
            findings.append(FileScanFinding(
                kind    = "prompt_injection",
                label   = f"Obfuscated Content (depth {depth})",
                excerpt = result_ob.decoded_extra[:80],
                line    = 0,
            ))
            if depth >= 2:
                injection_found = True
    except Exception as exc:
        log.debug("file_scan: obfuscation check skipped: %s", exc)

    risk  = _risk_level(findings, injection_found)
    safe  = risk in ("SAFE", "LOW")

    ms = (time.perf_counter() - t0) * 1000
    log.info(
        "file_scan tenant=%s file=%r size=%d risk=%s findings=%d ms=%.1f",
        tenant_id, filename, size_bytes, risk, len(findings), ms,
    )

    return FileScanResponse(
        filename       = filename,
        size_bytes     = size_bytes,
        safe           = safe,
        risk_level     = risk,
        findings       = findings,
        findings_count = len(findings),
        sanitized_text = sanitized,
        processing_ms  = round(ms, 2),
        truncated      = truncated,
    )


@router.get("/file/supported-types", summary="List supported file types")
async def supported_types() -> dict:
    return {
        "text": [".txt", ".md", ".csv", ".json", ".py", ".js", ".ts",
                 ".html", ".xml", ".yaml", ".env", ".sh", ".sql"],
        "pdf":  [".pdf"],
        "max_size_mb": _MAX_BYTES // 1024 // 1024,
        "office": [".docx", ".xlsx", ".xls"],
        "note": "Images and audio files are not supported.",
    }
